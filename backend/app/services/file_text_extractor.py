"""
文件文本抽取服务。

用于 JD 文件和批量简历文件的 PDF / Word / TXT 文本读取。
"""
from io import BytesIO
from pathlib import Path
import hashlib
import re
import unicodedata
from typing import Dict, Optional


SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}


def get_file_type(filename: str) -> str:
    ext = Path(filename or "").suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError("仅支持 PDF、DOC、DOCX、TXT 文件")
    return ext.lstrip(".").upper()


def extract_text_from_bytes(filename: str, content: bytes) -> str:
    ext = Path(filename or "").suffix.lower()
    if ext == ".pdf":
        text = _extract_pdf_text(content)
    elif ext == ".docx":
        text = _extract_docx_text(content)
    elif ext == ".doc":
        text = _extract_doc_text(content)
    elif ext == ".txt":
        text = _decode_text(content)
    else:
        raise ValueError("仅支持 PDF、DOC、DOCX、TXT 文件")

    text = _normalize_text(text)
    if len(text) < 10:
        raise ValueError("未能从文件中读取到有效文本")
    return text


def normalize_resume_text_for_hash(text: str) -> str:
    """生成简历内容指纹前的轻量归一化，降低导出格式差异带来的误判。"""
    normalized = unicodedata.normalize("NFKC", text or "").lower()
    normalized = re.sub(r"第\s*\d+\s*页\s*/?\s*共?\s*\d*\s*页?", " ", normalized)
    normalized = re.sub(r"\bpage\s*\d+\s*(of\s*\d+)?\b", " ", normalized)
    normalized = re.sub(r"[\s\r\n\t]+", "", normalized)
    normalized = re.sub(r"[·•●◆■□▪▫_—\-=\|/\\]+", "", normalized)
    return normalized.strip()


def compute_resume_text_hash(text: str) -> Optional[str]:
    normalized = normalize_resume_text_for_hash(text)
    if not normalized:
        return None
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def extract_resume_identity(text: str) -> Dict[str, str]:
    normalized = unicodedata.normalize("NFKC", text or "")
    phone_match = re.search(r"(?<!\d)(1[3-9]\d[\s\-]?\d{4}[\s\-]?\d{4})(?!\d)", normalized)
    email_match = re.search(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", normalized)
    phone = re.sub(r"\D", "", phone_match.group(1)) if phone_match else ""
    email = email_match.group(0).lower() if email_match else ""
    return {"phone": phone, "email": email}


def _extract_pdf_text(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("缺少 pypdf 依赖，请先安装后端依赖") from exc

    reader = PdfReader(BytesIO(content))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("缺少 python-docx 依赖，请先安装后端依赖") from exc

    document = Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells if cell.text]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_doc_text(content: bytes) -> str:
    # 老式 .doc 没有轻量、纯 Python 的稳定解析方案。这里先尽量从常见编码中抽取可读文本；
    # 如果抽取失败，接口会把原因记录到对应简历文件，HR 可转为 DOCX/PDF 后重试。
    candidates = [
        _decode_text(content, "utf-8"),
        _decode_text(content, "gb18030"),
        _decode_text(content, "utf-16le"),
        _decode_text(content, "latin-1"),
    ]
    return max(candidates, key=lambda text: len(_normalize_text(text)))


def _decode_text(content: bytes, preferred_encoding: Optional[str] = None) -> str:
    encodings = [preferred_encoding] if preferred_encoding else ["utf-8", "gb18030", "utf-16le"]
    for encoding in [item for item in encodings if item]:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[^\S\r\n]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
